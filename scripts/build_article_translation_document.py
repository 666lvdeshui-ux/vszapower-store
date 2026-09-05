"""Create paragraph-level translation input; keep Markdown syntax out of translation."""
import json,re
from pathlib import Path
from docx import Document
ROOT=Path(__file__).resolve().parents[1]/'content/battery-academy/translations'
source=json.loads((ROOT/'source-posts.json').read_text())['data']
entries=[];known={};refs={}
def phrase(s):
 if not s.strip() or not re.search(r'[A-Za-z\u4e00-\u9fff]',s):return {'literal':s}
 if s not in known:
  key='VSZSEG%05d'%len(entries);known[s]=key;entries.append({'key':key,'source':s})
 return {'ref':known[s]}
def field(s):
 # Preserve all Markdown delimiters, link URLs and formulas as immutable literals.
 parts=re.split(r'(\$\$[\s\S]*?\$\$|\$[^\n$]+\$|\]\([^\n)]+\)|https?://[^\s)]+|\n|\*\*|`|\||^#{1,6} |^[-*>] |^\d+\. |^[-: ]+$|\[|\])',s,flags=re.M)
 result=[]
 for part in parts:
  if not part:continue
  if part.startswith(('$$','$','](','http')) or not re.search(r'[A-Za-z\u4e00-\u9fff]',part):result.append({'literal':part});continue
  lead=part[:len(part)-len(part.lstrip())];trail=part[len(part.rstrip()):]
  if lead:result.append({'literal':lead})
  result.append(phrase(part.strip()))
  if trail:result.append({'literal':trail})
 return result
for p in source:
 refs[p['id']]={f:field(p.get(f,'') or '') for f in ['title','summary','category','content']}
 refs[p['id']]['tags']=[field(t) for t in p.get('tags',[])]
doc=Document()
for e in entries:
 doc.add_paragraph(e['key'],style='Heading 2');doc.add_paragraph(e['source'])
doc.save(ROOT/'translation-paragraphs.docx')
(ROOT/'paragraph-map.json').write_text(json.dumps({'entries':entries,'refs':refs},ensure_ascii=False,indent=2))
print({'phrases':len(entries),'characters':sum(len(e['source']) for e in entries)})
